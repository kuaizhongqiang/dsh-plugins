#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
parse_document.py - extract text and images from Word / Excel / PDF files.

Reads one office/PDF document and emits a JSON report on stdout:

    {
      "format": "docx" | "xlsx" | "pdf",
      "meta":    { ... per-format metadata ... },
      "text":    "...extracted text...",
      "images":  [ { "path": "abs/path/to/extracted.png", "mediaType": "image/png", "name": "..." } , ... ]
    }

Usage:
    python parse_document.py <input> [--out <image-out-dir>] [--max-images N]

Requirements (installed on the target machine):
    python-docx  (Word .docx)
    openpyxl     (Excel .xlsx)
    PyMuPDF      (PDF; import name `fitz`)

Exits 0 on success; the JSON report is printed to stdout (UTF-8). On any
failure a JSON object {"error": "..."} is printed and exit code is 1.
"""

import argparse
import io
import json
import os
import re
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

# --- force UTF-8 stdout so Chinese text survives Windows consoles -----------
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")


# ---------------------------------------------------------------- helpers ---
def emit(payload, exit_code=0):
    """Print the JSON payload and exit."""
    text = json.dumps(payload, ensure_ascii=False, indent=2)
    sys.stdout.write(text + "\n")
    sys.stdout.flush()
    sys.exit(exit_code)


def fail(message):
    emit({"error": message}, exit_code=1)


def safe_name(name, fallback="image"):
    """Sanitize a media entry name into a safe file stem."""
    stem = re.sub(r"[^\w\-.]", "_", name or "").strip("._")
    return stem or fallback


def save_bytes(data: bytes, out_dir: Path, stem: str, ext: str) -> Path:
    """Write bytes to out_dir/stem.ext, deduping collisions with a suffix."""
    ext = ext.lstrip(".").lower() or "bin"
    target = out_dir / f"{stem}.{ext}"
    n = 2
    while target.exists():
        target = out_dir / f"{stem}_{n}.{ext}"
        n += 1
    target.write_bytes(data)
    return target


MEDIA_EXT_TO_TYPE = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "webp": "image/webp",
    "bmp": "image/bmp",
    "tif": "image/tiff",
    "tiff": "image/tiff",
}


def media_type_for_ext(ext: str) -> str:
    return MEDIA_EXT_TO_TYPE.get(ext.lstrip(".").lower(), "application/octet-stream")


def sniff_image_type(data: bytes) -> str:
    """Best-effort media type detection from magic bytes."""
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if data[:2] == b"\xff\xd8":
        return "image/jpeg"
    if data[:6] in (b"GIF87a", b"GIF89a"):
        return "image/gif"
    if data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "image/webp"
    if data[:2] == b"BM":
        return "image/bmp"
    if data[:4] in (b"II*\x00", b"MM\x00*"):
        return "image/tiff"
    return "application/octet-stream"


def media_ext_for_type(media_type: str) -> str:
    return {
        "image/png": "png",
        "image/jpeg": "jpg",
        "image/gif": "gif",
        "image/webp": "webp",
        "image/bmp": "bmp",
        "image/tiff": "tiff",
    }.get(media_type, "bin")


# ---------------------------------------------------------------- docx -------
def extract_docx(input_path: Path, out_dir: Path, max_images: int):
    import docx  # python-docx

    document = docx.Document(str(input_path))

    parts = []
    for para in document.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    # tables: render each cell in row-major order with row separators
    for ti, table in enumerate(document.tables):
        parts.append(f"[表格 {ti + 1}]")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
        parts.append(f"[/表格 {ti + 1}]")

    text = "\n".join(parts)

    # images live inside word/media/ of the OOXML zip
    images = []
    try:
        with zipfile.ZipFile(input_path) as zf:
            media_names = sorted(
                n for n in zf.namelist()
                if n.startswith("word/media/") and not n.endswith("/")
            )
            if max_images is not None and max_images >= 0:
                media_names = media_names[:max_images]
            for name in media_names:
                data = zf.read(name)
                media_type = media_type_for_ext(Path(name).suffix)
                if media_type == "application/octet-stream":
                    media_type = sniff_image_type(data)
                stem = safe_name(Path(name).stem, "docx_image")
                ext = media_ext_for_type(media_type)
                if ext == "bin":
                    continue  # skip non-image media entries
                path = save_bytes(data, out_dir, stem, ext)
                images.append({
                    "path": str(path),
                    "mediaType": media_type,
                    "name": Path(name).name,
                })
    except zipfile.BadZipFile:
        pass  # no media readable; text still valid

    meta = {
        "format": "docx",
        "fileName": input_path.name,
        "paragraphCount": len(document.paragraphs),
        "tableCount": len(document.tables),
    }
    return {"format": "docx", "meta": meta, "text": text, "images": images}


# ---------------------------------------------------------------- xlsx -------
def extract_xlsx(input_path: Path, out_dir: Path, max_images: int):
    import openpyxl

    wb = openpyxl.load_workbook(str(input_path), data_only=True, read_only=True)

    parts = []
    for sheet in wb.worksheets:
        parts.append(f"=== 工作表: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            cells = []
            for value in row:
                if value is None:
                    continue
                if isinstance(value, float) and value.is_integer():
                    value = int(value)
                cells.append(str(value))
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)

    # embedded images in xl/media/ of the zip
    images = []
    try:
        with zipfile.ZipFile(input_path) as zf:
            media_names = sorted(
                n for n in zf.namelist()
                if n.startswith("xl/media/") and not n.endswith("/")
            )
            if max_images is not None and max_images >= 0:
                media_names = media_names[:max_images]
            for name in media_names:
                data = zf.read(name)
                media_type = media_type_for_ext(Path(name).suffix)
                if media_type == "application/octet-stream":
                    media_type = sniff_image_type(data)
                ext = media_ext_for_type(media_type)
                if ext == "bin":
                    continue
                stem = safe_name(Path(name).stem, "xlsx_image")
                path = save_bytes(data, out_dir, stem, ext)
                images.append({
                    "path": str(path),
                    "mediaType": media_type,
                    "name": Path(name).name,
                })
    except zipfile.BadZipFile:
        pass

    meta = {
        "format": "xlsx",
        "fileName": input_path.name,
        "sheetCount": len(wb.worksheets),
        "sheetNames": [s.title for s in wb.worksheets],
    }
    return {"format": "xlsx", "meta": meta, "text": text, "images": images}


# ---------------------------------------------------------------- pdf --------
def extract_pdf(input_path: Path, out_dir: Path, max_images: int):
    import fitz  # PyMuPDF

    doc = fitz.open(str(input_path))

    parts = []
    for page_no, page in enumerate(doc):
        t = page.get_text("text").strip()
        if t:
            parts.append(f"--- 第 {page_no + 1} 页 ---\n{t}")

    # embedded raster images
    images = []
    seen = set()
    for page_no, page in enumerate(doc):
        for img in page.get_images(full=True):
            xref = img[0]
            if xref in seen:
                continue
            seen.add(xref)
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.n - pix.alpha > 3:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                png = pix.tobytes("png")
                ext = "png"
                media_type = "image/png"
            except Exception:
                try:
                    raw = doc.extract_image(xref)
                    media_type = raw.get("ext", "png")
                    media_type = {
                        "png": "image/png",
                        "jpeg": "image/jpeg",
                        "jpg": "image/jpeg",
                        "gif": "image/gif",
                        "webp": "image/webp",
                        "bmp": "image/bmp",
                    }.get(media_type, "image/png")
                    png = raw.get("image", b"")
                    ext = media_ext_for_type(media_type)
                except Exception:
                    continue
            if not png:
                continue
            if max_images is not None and max_images >= 0 and len(images) >= max_images:
                break
            path = save_bytes(png, out_dir, f"pdf_p{page_no + 1}_img{xref}", ext)
            images.append({
                "path": str(path),
                "mediaType": media_type,
                "name": f"page{page_no + 1}_img{xref}.{ext}",
            })
        if max_images is not None and max_images >= 0 and len(images) >= max_images:
            break

    meta = {
        "format": "pdf",
        "fileName": input_path.name,
        "pageCount": doc.page_count,
    }
    return {"format": "pdf", "meta": meta, "text": "\n".join(parts), "images": images}


# ---------------------------------------------------------------- main -------
def main():
    parser = argparse.ArgumentParser(description="Extract text and images from docx/xlsx/pdf.")
    parser.add_argument("input", help="Path to the input document")
    parser.add_argument("--out", default=None, help="Directory for extracted images (default: temp)")
    parser.add_argument("--max-images", type=int, default=-1, help="Cap on extracted images (-1 = unlimited)")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.is_file():
        fail(f"input file not found: {args.input}")

    ext = input_path.suffix.lower()
    out_dir = Path(args.out) if args.out else Path(tempfile.mkdtemp(prefix="dsh_doc_"))
    out_dir.mkdir(parents=True, exist_ok=True)

    max_images = args.max_images if args.max_images is not None and args.max_images >= 0 else None

    try:
        if ext in (".docx", ".docm"):
            result = extract_docx(input_path, out_dir, max_images)
        elif ext == ".xlsx":
            result = extract_xlsx(input_path, out_dir, max_images)
        elif ext == ".pdf":
            result = extract_pdf(input_path, out_dir, max_images)
        else:
            fail(f"unsupported format: {ext} (supported: .docx .docm .xlsx .pdf)")
    except ImportError as exc:
        missing = exc.name or str(exc)
        fail(f"missing python library for {ext} parsing: {missing}. "
             f"Install it, e.g. `pip install {missing}`")
    except Exception as exc:  # noqa: BLE001 - report anything to the caller
        fail(f"failed to parse {ext}: {type(exc).__name__}: {exc}")

    emit(result)


if __name__ == "__main__":
    main()
